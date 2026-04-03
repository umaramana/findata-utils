"""
merge_gvk.py
Merges two GVK verse files (TXT + DOCX) by matching verse numbers.
Output retains file2 (DOCX) formatting — file2 is used as the document template.

Per-verse output order:
  [chapter title / heading — deep-copied from file2 with original style]
  N.\t[TXT verse lines 1–4, soft returns between lines]
  [TXT remaining lines — padachedam, arumpadhavurai, pozhippurai, etc.]
  [blank line — visual separator between file1 and file2 content]
  [DOCX lines — number stripped]
  [blank line — end of verse]

Usage:
  python merge_gvk.py                   # uses default files in gvk/ root
  python merge_gvk.py trialphase1and2   # finds *file1.txt + *file2.docx in subfolder
  python merge_gvk.py --verify-only     # check verse counts only, no output written
  python merge_gvk.py trialphase1and2 --verify-only
"""

import re
import sys
import shutil
from collections import defaultdict
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FOLDER = Path(__file__).parent

# Default file paths (used when no folder arg is given)
_DEFAULT_TXT = FOLDER / "gvk sample 1-62.txt"
_DEFAULT_DOC = FOLDER / "sample 1-62.docx"
_DEFAULT_OUT = FOLDER / "gvk_merged.docx"

VERSE_PAT     = re.compile(r'^(\d+)\.\s*(.+)')
MIN_VERSE_LEN = 40   # section headers are shorter; real verses are longer

# Styles treated as chapter-level headers to preserve between verses.
# Heading 3 and below are verse-level content (e.g. விளக்கக் குறிப்பு:) — left as verse content.
CHAPTER_HEADER_STYLES = {'Title', 'Heading 1', 'Heading 2'}


# ── Path resolution ────────────────────────────────────────────────────────────


def resolve_paths():
    """
    No folder arg → default files in gvk/ root.
    Folder arg   → look for *file1.txt and *file2.docx in that subfolder;
                   output is placed in the same subfolder as *merged.docx.
    """
    folder_arg = next((a for a in sys.argv[1:] if not a.startswith("--")), None)

    if folder_arg is None:
        return _DEFAULT_TXT, _DEFAULT_DOC, _DEFAULT_OUT

    folder = FOLDER / folder_arg
    if not folder.is_dir():
        sys.exit(f"ERROR: Folder not found: {folder}")

    txt_files = list(folder.glob("*file1.txt"))
    doc_files = list(folder.glob("*file2.docx"))

    if not txt_files:
        sys.exit(f"ERROR: No *file1.txt found in {folder}")
    if not doc_files:
        sys.exit(f"ERROR: No *file2.docx found in {folder}")

    txt = txt_files[0]
    doc = doc_files[0]
    stem = doc.stem.replace("file2", "merged")
    out  = folder / f"{stem}.docx"
    return txt, doc, out


# ── Helpers ────────────────────────────────────────────────────────────────────

def _gaps(nums):
    """Return missing integers within min..max of the given set."""
    if not nums:
        return []
    s = set(nums)
    return [n for n in range(min(s), max(s) + 1) if n not in s]


def _strip_trailing_blank(lines):
    while lines and not lines[-1].strip():
        lines.pop()
    return lines


def _is_section_header_txt(line):
    """
    TXT section headers:
    - Tab after number: '1.\tகுரு வணக்கம்'  (chapter header in body)
    - Page number suffix: '63. Chapter\t\t128'  (table of contents entry)
    """
    if re.match(r'^\d+\.\t', line):
        return True
    if re.search(r'\t\d+$', line):
        return True
    return False


# ── Parsers ────────────────────────────────────────────────────────────────────

def parse_txt(path):
    """
    Return {verse_num: [lines]}.
    Lines are plain text only — verse number is NOT stored (it is the dict key).
    Skips section headers (tab after number) and duplicate verse numbers.
    """
    verses = {}
    current_num   = None
    current_lines = []

    # Detect encoding from BOM; fall back to UTF-8
    with open(path, 'rb') as _fb:
        _bom = _fb.read(2)
    _enc = 'utf-16' if _bom in (b'\xff\xfe', b'\xfe\xff') else 'utf-8'

    with open(path, encoding=_enc) as f:
        for raw in f:
            line = raw.rstrip("\n")

            if _is_section_header_txt(line):
                if current_num is not None:
                    verses[current_num] = _strip_trailing_blank(current_lines)
                current_num   = None
                current_lines = []
                continue

            m = VERSE_PAT.match(line)
            if m:
                n = int(m.group(1))
                if n in verses:
                    # Duplicate number = sub-section header — flush and skip
                    if current_num is not None:
                        verses[current_num] = _strip_trailing_blank(current_lines)
                    current_num   = None
                    current_lines = []
                    continue
                if current_num is not None:
                    verses[current_num] = _strip_trailing_blank(current_lines)
                current_num   = n
                current_lines = [m.group(2).strip()]   # text only, no "N." prefix
            else:
                if current_num is not None:
                    current_lines.append(line)

    if current_num is not None:
        verses[current_num] = _strip_trailing_blank(current_lines)

    return verses


def parse_docx(path):
    """
    Return (verses, chapter_headers) where:
      verses          = {verse_num: [para_elements]}  ← raw XML elements, NOT text
      chapter_headers = [(p_element, next_verse_num), ...]

    Paragraph elements are stored so they can be deep-copied into the output,
    preserving all run-level formatting (red font, yellow highlight, bold, etc.).
    Blank paragraphs are filtered out at collection time.

    Chapter headers (Title, Heading 1, Heading 2) are pulled out and associated
    with the verse number that follows them.
    Verse-level content (e.g. Heading 3 'விளக்கக் குறிப்பு:') stays in the
    verse's element list.
    """
    doc = Document(path)
    verses          = {}
    chapter_headers = []        # [(p_element, next_verse_num)]
    current_num     = None
    current_elems   = []
    pending_hdrs    = []        # chapter headers waiting for the next verse

    for para in doc.paragraphs:
        # Use full XML text extraction so inline sdt content is included
        full_text  = ''.join(t.text or '' for t in para._element.iter(qn('w:t')))
        style_name = para.style.name

        m = VERSE_PAT.match(full_text)
        if m:
            content = m.group(2).strip()
            if len(content) < MIN_VERSE_LEN:
                # Short verse-pattern line = chapter/section heading
                pending_hdrs.append(para._element)
                continue
            n = int(m.group(1))
            if current_num is not None:
                verses[current_num] = current_elems
            # Flush pending headers — they belong before verse n
            for elem in pending_hdrs:
                chapter_headers.append((elem, n))
            pending_hdrs  = []
            current_num   = n
            current_elems = [para._element]   # first elem = verse para (num stripped on write)
        else:
            if style_name in CHAPTER_HEADER_STYLES and full_text.strip():
                # Non-numbered chapter title (e.g. 'பாயிரம்')
                pending_hdrs.append(para._element)
            elif current_num is not None and full_text.strip():
                # Verse content — keep non-blank paragraphs only
                current_elems.append(para._element)

    if current_num is not None:
        verses[current_num] = current_elems

    return verses, chapter_headers


# ── Verification ───────────────────────────────────────────────────────────────

def verify(txt_verses, docx_verses):
    txt_keys  = set(txt_verses.keys())
    docx_keys = set(docx_verses.keys())
    only_txt  = sorted(txt_keys - docx_keys)
    only_docx = sorted(docx_keys - txt_keys)
    common    = sorted(txt_keys & docx_keys)
    txt_gaps  = _gaps(txt_keys)
    docx_gaps = _gaps(docx_keys)

    print(f"TXT  verses detected : {len(txt_keys)}  ({min(txt_keys)}-{max(txt_keys)})")
    if txt_gaps:
        print(f"  WARNING — TXT  gaps : {txt_gaps}")

    print(f"DOCX verses detected : {len(docx_keys)}  ({min(docx_keys)}-{max(docx_keys)})")
    if docx_gaps:
        print(f"  WARNING — DOCX gaps : {docx_gaps}")

    print(f"Matched              : {len(common)}")

    if only_txt:
        print(f"WARNING — in TXT only  : {only_txt}")
    if only_docx:
        print(f"WARNING — in DOCX only : {only_docx}")

    ok = not any([only_txt, only_docx, txt_gaps, docx_gaps])
    if ok:
        print("OK — verse numbers match perfectly.")
    return ok


def verify_output(out_path, expected_nums):
    """Check merged DOCX for verse contiguity against expected verse numbers."""
    doc   = Document(out_path)
    found = set()
    for para in doc.paragraphs:
        m = re.match(r'^(\d+)\.\t', para.text)
        if m:
            found.add(int(m.group(1)))

    print(f"\nOutput check:")
    if not found:
        print("  ERROR — no verses found in output!")
        return
    print(f"  Verses in output : {len(found)}  ({min(found)}-{max(found)})")

    gaps    = _gaps(found)
    missing = sorted(set(expected_nums) - found)
    extra   = sorted(found - set(expected_nums))

    if gaps:
        print(f"  WARNING — gaps in output   : {gaps}")
    if missing:
        print(f"  WARNING — missing in output: {missing}")
    if extra:
        print(f"  WARNING — extra in output  : {extra}")
    if not gaps and not missing and not extra:
        print("  OK — output matches input exactly.")


# ── Formatting helpers ─────────────────────────────────────────────────────────

def _build_rpr():
    """
    Build a run properties (rPr) element matching file2's normal paragraph runs:
      - Font: Latha (ascii, cs, eastAsia, hAnsi)
      - Size: 10pt (sz=20 half-points, szCs=20 for complex script)
      - RTL: off
    """
    rpr    = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Latha')
    rFonts.set(qn('w:cs'),       'Latha')
    rFonts.set(qn('w:eastAsia'), 'Latha')
    rFonts.set(qn('w:hAnsi'),    'Latha')
    rpr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rpr.append(sz)

    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '20')
    rpr.append(szCs)

    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '0')
    rpr.append(rtl)

    return rpr


def _apply_fmt(run):
    """Stamp file2's normal run formatting (Latha 10pt) onto a run."""
    r = run._r
    existing = r.find(qn('w:rPr'))
    if existing is not None:
        r.remove(existing)
    r.insert(0, _build_rpr())


def _set_spacing(para):
    """Apply 1.5 line spacing to a paragraph (safety fallback)."""
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def _add_soft_break(para):
    """Insert a soft return (Shift+Enter) into a paragraph."""
    run = para.add_run()
    br  = OxmlElement('w:br')
    run._r.append(br)
    _apply_fmt(run)


def _insert_elem(out, elem):
    """Deep-copy a paragraph element from file2 and append it to the output body."""
    new_elem = deepcopy(elem)
    body     = out.element.body
    sectPr   = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(new_elem)
    else:
        body.append(new_elem)


def _insert_elem_strip_num(out, elem):
    """
    Deep-copy a file2 verse paragraph, strip the leading verse number and tab
    from the first run, then insert into output.

    In file2 the verse number is stored as separate XML nodes inside the first w:r:
      <w:t>N.</w:t>  <w:tab/>  <w:t>actual text...</w:t>
    We remove the <w:t>N.</w:t> and the <w:tab/> that follows it.
    All other run formatting (red font, yellow highlight, etc.) is preserved.
    """
    new_elem  = deepcopy(elem)
    first_r   = new_elem.find('.//' + qn('w:r'))
    if first_r is not None:
        to_remove = []
        found_num = False
        for child in list(first_r):
            if (not found_num
                    and child.tag == qn('w:t')
                    and child.text
                    and re.match(r'^\d+\.$', child.text.strip())):
                to_remove.append(child)
                found_num = True
            elif found_num and child.tag == qn('w:tab'):
                to_remove.append(child)
                break
        for child in to_remove:
            first_r.remove(child)

    body   = out.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(new_elem)
    else:
        body.append(new_elem)


# ── Verify helpers ────────────────────────────────────────────────────────────

def _is_latha_para(para):
    """True if any run in para carries Latha font — indicates file1-sourced content."""
    for r in para._element.findall('.//' + qn('w:r')):
        rpr = r.find(qn('w:rPr'))
        if rpr is None:
            continue
        rf = rpr.find(qn('w:rFonts'))
        if rf is not None and rf.get(qn('w:ascii')) == 'Latha':
            return True
    return False


def write_verify_report(txt_verses, out_path):
    """
    Compare file1 expected content against file1-sourced paragraphs in the
    merged output, verse by verse.  Writes verify_report.txt to the same folder.
    """
    doc = Document(out_path)

    # Group output paragraphs by verse number
    verse_paras = {}
    current = None
    for para in doc.paragraphs:
        m = re.match(r'^(\d+)\.\t', para.text)
        if m:
            current = int(m.group(1))
            verse_paras[current] = [para]
        elif current is not None:
            verse_paras[current].append(para)

    report_lines = []
    mismatches   = 0

    for num in sorted(set(txt_verses.keys()) | set(verse_paras.keys())):
        report_lines.append(f"=== Verse {num} ===")

        # Build expected lines from txt_verses
        txt_block   = txt_verses.get(num, [])
        verse_lines = txt_block[:4]
        rest_lines  = txt_block[4:]
        expected = []
        if verse_lines:
            expected.append(f"{num}.\t{verse_lines[0]}")
            for l in verse_lines[1:]:
                expected.append(f"\t{l}")
        for l in rest_lines:
            expected.append(l)

        # Actual: file1-sourced (Latha), non-blank paragraphs for this verse
        actual = [
            p.text for p in verse_paras.get(num, [])
            if _is_latha_para(p) and p.text.strip()
        ]

        report_lines.append("[FILE1]")
        report_lines.extend(f"  {l}" for l in expected)
        report_lines.append("[OUTPUT]")
        report_lines.extend(f"  {l}" for l in actual)

        if expected == actual:
            report_lines.append("MATCH")
        else:
            report_lines.append("MISMATCH")
            mismatches += 1
        report_lines.append("---")

    report_path = out_path.parent / "verify_report.txt"
    report_path.write_text("\n".join(report_lines), encoding="utf-8")
    print(f"\nVerify report: {report_path}")
    print(f"  {len(txt_verses)} verses checked, {mismatches} mismatch(es)")


# ── Merge & write ──────────────────────────────────────────────────────────────

def write_merged(txt_verses, docx_verses, chapter_headers, doc_file, out_path):
    """
    Copy file2 as the base document so its styles and fonts are inherited,
    then clear its content and write merged verses.
    Chapter headers are deep-copied from file2 to preserve exact styling.
    """
    shutil.copy(doc_file, out_path)
    out = Document(out_path)

    # Clear all body content (paragraphs, sdts, bookmarks) — preserve only sectPr and styles
    body = out.element.body
    sectPr_tag = qn('w:sectPr')
    for child in list(body):
        if child.tag != sectPr_tag:
            body.remove(child)

    # Build lookup: verse_num → [header elements to insert before it]
    headers_before = defaultdict(list)
    for elem, verse_num in chapter_headers:
        headers_before[verse_num].append(elem)

    all_nums = sorted(set(txt_verses.keys()) | set(docx_verses.keys()))

    for num in all_nums:
        # ── Chapter headers before this verse (deep-copied from file2) ─────
        for elem in headers_before.get(num, []):
            _insert_elem(out, elem)

        txt_block  = [l for l in txt_verses.get(num, []) if l.strip()]
        docx_block = docx_verses.get(num, [])

        verse_lines = txt_block[:4]
        rest_lines  = txt_block[4:]

        # ── Verse lines: each on its own paragraph, lines 2–4 indented with tab ──
        if verse_lines:
            p = out.add_paragraph(style='Normal')
            _set_spacing(p)
            run = p.add_run(f"{num}.\t{verse_lines[0]}")
            _apply_fmt(run)
            for line in verse_lines[1:]:
                p = out.add_paragraph(style='Normal')
                _set_spacing(p)
                run = p.add_run(f"\t{line}")
                _apply_fmt(run)
            out.add_paragraph(style='Normal')   # blank after verse block

        # ── Remaining TXT lines (padachedam, arumpadhavurai, pozhippurai…) ─
        # Each section followed by a blank line (2 hard returns between sections)
        for line in rest_lines:
            p = out.add_paragraph(style='Normal')
            _set_spacing(p)
            run = p.add_run(line)
            _apply_fmt(run)
            out.add_paragraph(style='Normal')   # blank after each section

        # ── DOCX paragraphs — deep-copied to preserve all run formatting ──
        # (red font, yellow highlight, bold, etc.)
        for i, elem in enumerate(docx_block):
            if i == 0:
                _insert_elem_strip_num(out, elem)   # strip "N.\t" from verse para
            else:
                _insert_elem(out, elem)

        # ── Hard return after each verse ────────────────────────────────────
        out.add_paragraph(style='Normal')

    out.save(out_path)
    print(f"\nSaved: {out_path}")


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    txt_file, doc_file, out_file = resolve_paths()

    print(f"TXT  : {txt_file.name}")
    print(f"DOCX : {doc_file.name}")
    print(f"OUT  : {out_file.name}")
    print()

    print("Parsing TXT ...")
    txt_verses = parse_txt(txt_file)

    print("Parsing DOCX ...")
    docx_verses, chapter_headers = parse_docx(doc_file)

    print()
    ok = verify(txt_verses, docx_verses)

    if "--verify-only" in sys.argv:
        return

    if not ok:
        ans = input("\nMismatches found. Continue anyway? [y/N] ").strip().lower()
        if ans != "y":
            print("Aborted.")
            return

    print("\nWriting merged output ...")
    write_merged(txt_verses, docx_verses, chapter_headers, doc_file=doc_file, out_path=out_file)

    expected_nums = set(txt_verses.keys()) & set(docx_verses.keys())
    verify_output(out_file, expected_nums)

    if "--verify-report" in sys.argv:
        write_verify_report(txt_verses, out_file)


if __name__ == "__main__":
    main()
